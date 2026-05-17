"""
Docx Master — Document Scanner
Usage: python scan_docx.py <input.docx> [output_template.json]

Reads a .docx and extracts its structure as a JSON template suitable for
generate_docx.py.  Auto-installs python-docx if missing.
"""

import json
import subprocess
import sys
import os
from pathlib import Path

# ── Auto-install python-docx if missing ──────────────────────────
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[docx-master] python-docx not found, installing...")
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    print("[docx-master] python-docx installed OK")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn


# ── Helpers ───────────────────────────────────────────────────────

HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "heading 1", "heading 2", "heading 3", "heading 4",
    "1. Heading 1", "2. Heading 2", "3. Heading 3",   # WPS variants
}
TITLE_STYLES = {"Title", "title", "Document Title"}
BULLET_STYLES = {"List Bullet", "List Paragraph", "List Bullet 2"}

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT:   "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT:  "right",
}


def _get_heading_level(style_name):
    """Extract heading level from style name. Returns 1-4 or None."""
    for i in range(1, 5):
        if str(i) in style_name:
            return i
    return None


def _get_alignment(para):
    """Return string alignment of a paragraph."""
    if para.alignment is None:
        return None
    return ALIGN_MAP.get(para.alignment)


def _extract_paragraph(para):
    """Classify and extract a single paragraph."""
    style_name = para.style.name if para.style else ""

    # Skip empty paragraphs (but not if they're spacing between sections)
    text = para.text.strip()
    if not text and style_name not in HEADING_STYLES:
        return None

    # Heading detection by style
    level = _get_heading_level(style_name)
    if level:
        return {"type": "heading", "text": text, "level": level}

    # Title detection
    if style_name in TITLE_STYLES:
        return {"type": "heading", "text": text, "level": 0}

    # Bullet detection
    if style_name in BULLET_STYLES or para.text.strip().startswith("•"):
        clean = text.lstrip("• ").lstrip("- ").strip()
        return {"type": "bullet", "text": clean}

    # Normal paragraph — extract formatting from first run
    item = {"type": "paragraph", "text": text}
    if para.runs:
        r0 = para.runs[0]
        if r0.bold:
            item["bold"] = True
        if r0.italic:
            item["italic"] = True
        if r0.font.size:
            item["sz"] = str(int(r0.font.size.pt * 2))   # pt → half-pts
        if r0.font.color and r0.font.color.rgb:
            item["color"] = str(r0.font.color.rgb)
    align = _get_alignment(para)
    if align:
        item["align"] = align
    return item


def _extract_table(table):
    """Extract a table as JSON."""
    headers = []
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if i == 0:
            headers = cells
        else:
            rows.append(cells)
    if not headers:
        return None
    return {"type": "table", "headers": headers, "rows": rows}


def _is_title_paragraph(para):
    """Heuristic: a large, bold, centered paragraph without a heading style
    likely functions as the document title."""
    if not para.runs or not para.text.strip():
        return False
    r0 = para.runs[0]
    align = _get_alignment(para)
    if align != "center":
        return False
    if r0.font.size and r0.font.size.pt >= 16:
        return True
    return False


# ── Main scanner ──────────────────────────────────────────────────

def scan_docx(input_path):
    """Read a .docx and return a JSON-serializable content dict."""
    doc = Document(input_path)

    result = {"title": "", "sections": []}
    current_section = {"heading": "", "level": 1, "content": []}
    body_elements = []   # collect all body children in order

    # Collect all body children in document order
    body = doc.element.body

    para_index = 0
    table_index = 0
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_index < len(doc.paragraphs):
                para = doc.paragraphs[para_index]
                body_elements.append(("p", para))
                para_index += 1
        elif tag == "tbl":
            if table_index < len(doc.tables):
                tbl = doc.tables[table_index]
                body_elements.append(("tbl", tbl))
                table_index += 1

    # First pass: detect title (first centered large text or Title style)
    for tag, obj in body_elements:
        if tag == "p":
            style_name = obj.style.name if obj.style else ""
            if style_name in TITLE_STYLES or _is_title_paragraph(obj):
                result["title"] = obj.text.strip()
                break

    # Second pass: build sections
    for tag, obj in body_elements:
        if tag == "p":
            style_name = obj.style.name if obj.style else ""
            text = obj.text.strip()

            # Skip the title if we already captured it
            if text == result["title"] and (
                style_name in TITLE_STYLES or _is_title_paragraph(obj)
            ):
                continue

            # New heading → flush current section, start new one
            if _get_heading_level(style_name):
                # Flush previous section if it has content
                if current_section["heading"] or current_section["content"]:
                    result["sections"].append(current_section)

                level = _get_heading_level(style_name) or 1
                current_section = {
                    "heading": text,
                    "level": level,
                    "content": [],
                }
                continue

            extracted = _extract_paragraph(obj)
            if extracted:
                current_section["content"].append(extracted)

        elif tag == "tbl":
            extracted = _extract_table(obj)
            if extracted:
                current_section["content"].append(extracted)

    # Flush last section
    if current_section["heading"] or current_section["content"]:
        result["sections"].append(current_section)

    # If there's no title found and first element is a heading, use it
    if not result["title"] and result["sections"]:
        first = result["sections"][0]
        if first["heading"] and not first["content"]:
            result["title"] = first["heading"]
            result["sections"].pop(0)

    return result


def main():
    if len(sys.argv) < 2:
        print("Usage: python scan_docx.py <input.docx> [output.json]")
        print("  If output.json is omitted, prints JSON to stdout.")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.isfile(input_path):
        print(f"ERROR: file not found: {input_path}")
        sys.exit(1)

    result = scan_docx(input_path)

    json_str = json.dumps(result, ensure_ascii=False, indent=2)

    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(json_str)
        print(f"OK: {out_path}")
    else:
        print(json_str)


if __name__ == "__main__":
    main()
