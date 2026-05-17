"""
Docx Master — python-docx generator
Usage: python generate_docx.py <output.docx> <content.json>

Auto-installs python-docx if missing, then generates a .docx from JSON.
"""

import json
import subprocess
import sys
import os
from pathlib import Path

# ── Auto-install python-docx if missing ──────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("[docx-master] python-docx not found, installing...")
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    print("[docx-master] python-docx installed OK")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

# ── Configuration ─────────────────────────────────────────────────
FONT_EAST_ASIA = "宋体"
FONT_ASCII = "Times New Roman"

ALIGN_MAP = {
    "left":   WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right":  WD_ALIGN_PARAGRAPH.RIGHT,
}


def setup_styles(doc):
    """Chinese-friendly defaults: 宋体 body, sized headings 1-4."""
    style = doc.styles["Normal"]
    style.font.name = FONT_ASCII
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    for i, sz in {1: 18, 2: 14, 3: 13, 4: 12}.items():
        hname = f"Heading {i}"
        if hname in [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]:
            hs = doc.styles[hname]
        else:
            hs = doc.styles.add_style(hname, WD_STYLE_TYPE.PARAGRAPH)
            hs.base_style = doc.styles["Normal"]
        hs.font.name = FONT_ASCII
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        hs.font.bold = True
        hs.font.size = Pt(sz)


def _font_run(run):
    run.font.name = FONT_ASCII
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


# ── Content builders ──────────────────────────────────────────────

def add_paragraph(doc, text, bold=False, italic=False, align=None,
                  sz=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    _font_run(run)
    if sz:
        run.font.size = Pt(int(sz) / 2)   # half-pts → pt
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align in ALIGN_MAP:
        p.alignment = ALIGN_MAP[align]


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    _font_run(run)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _font_run(run)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, txt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        _font_run(run)
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r, row_data in enumerate(rows):
        for c, txt in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(txt))
            _font_run(run)
            run.font.size = Pt(10.5)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Emu(w)

    doc.add_paragraph()   # spacer


def add_image(doc, image_path, width=None, height=None, caption=None):
    if not os.path.isfile(image_path):
        print(f"[docx-master] WARNING: image not found: {image_path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    kwargs = {}
    if width:
        kwargs["width"] = Emu(width)
    if height:
        kwargs["height"] = Emu(height)
    if not kwargs:
        kwargs["width"] = Inches(5.5)

    try:
        run.add_picture(image_path, **kwargs)
    except Exception as exc:
        print(f"[docx-master] WARNING: cannot embed {image_path}: {exc}")
        return

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.size = Pt(10)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _font_run(cr)


# ── Main ──────────────────────────────────────────────────────────

CONTENT_HANDLERS = {
    "paragraph": lambda doc, item: add_paragraph(
        doc, item["text"],
        bold=item.get("bold", False),
        italic=item.get("italic", False),
        align=item.get("align"),
        sz=item.get("sz"),
        color=item.get("color"),
    ),
    "heading": lambda doc, item: add_heading(
        doc, item["text"], item.get("level", 2),
    ),
    "bullet": lambda doc, item: add_bullet(doc, item["text"]),
    "table": lambda doc, item: add_table(
        doc, item["headers"], item["rows"], item.get("colWidths"),
    ),
    "image": lambda doc, item: add_image(
        doc, item["path"],
        item.get("width"), item.get("height"), item.get("caption"),
    ),
}


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <output.docx> <content.json>")
        sys.exit(1)

    out_path = sys.argv[1]
    with open(sys.argv[2], "r", encoding="utf-8") as fh:
        content = json.load(fh)

    doc = Document()
    setup_styles(doc)

    # Document title
    if content.get("title"):
        t = doc.add_heading(content["title"], level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            _font_run(run)

    # Sections
    img_count = 0
    for sec in content.get("sections", []):
        if sec.get("heading"):
            add_heading(doc, sec["heading"], sec.get("level", 1))
        for item in sec.get("content", []):
            handler = CONTENT_HANDLERS.get(item.get("type"))
            if handler:
                handler(doc, item)
                if item.get("type") == "image":
                    img_count += 1

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    print(f"OK: {out_path} ({os.path.getsize(out_path)} bytes, {img_count} images)")


if __name__ == "__main__":
    main()
